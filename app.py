
import streamlit as st
import PyPDF2
import nltk
from collections import Counter
from docx import Document
import difflib  # For calculating similarity in plagiarism check
from dotenv import load_dotenv
load_dotenv()
import base64
import os
from PIL import Image
import pdf2image
import google.generativeai as genai
from io import BytesIO
from fpdf import FPDF

nltk.download('punkt')
nltk.download('stopwords')
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize


# Configure Google Generative AI
genai.configure(api_key=("AIzaSyAHrZuX0mzeKQUTtUNX1zv3dNO-m56T_nU"))


def get_gemini_response(input, pdf_content, prompt):
    model = genai.GenerativeModel('gemini-1.5-flash')
    response = model.generate_content([input, pdf_content[0], prompt])
    return response.text

def get_gemini_response1(input_prompt):
    model = genai.GenerativeModel('gemini-1.5-flash')
    response = model.generate_content(input_prompt)
    return response.text

def extract_text_from_pdf(pdf_file):
    reader = PyPDF2.PdfReader(pdf_file)
    text = ''
    for page_num in range(len(reader.pages)):
        page = reader.pages[page_num]
        text += page.extract_text()
    return text


def input_pdf_setup(pdf_file):
    return [extract_text_from_pdf(pdf_file)]

def extract_skills(text):
    skill_set = {'python', 'java', 'data analysis', 'project management', 'machine learning', 'communication', 'sql'}
    tokens = word_tokenize(text.lower())
    cleaned_tokens = [word for word in tokens if word.isalpha() and word not in stopwords.words('english')]
    relevant_skills = skill_set.intersection(cleaned_tokens)
    return relevant_skills

def generate_pdf(content):
    # Function to generate a PDF from the given resume content
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    # Split content into lines and add each line to the PDF
    for line in content.split("\n"):
        pdf.cell(200, 10, txt=line, ln=True)

    # Save the PDF to a BytesIO object
    pdf_output = BytesIO()
    pdf.output(pdf_output)
    pdf_output.seek(0)

    return pdf_output

def generate_pdf(content):
    # Function to generate a PDF from the given resume content
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    # Replace unsupported Unicode characters
    content = content.replace("\u2013", "-")  # Replace en-dash with a hyphen
    content = content.replace("\u2014", "--")  # Replace em-dash with double hyphen
    content = content.replace("\u2022", "*")  # Replace bullet points with asterisks

    # Split content into lines and add each line to the PDF
    for line in content.split("\n"):
        try:
            pdf.cell(200, 10, txt=line.encode('latin-1', 'replace').decode('latin-1'), ln=True)
        except UnicodeEncodeError as e:
            st.error(f"Error encoding line: {line} - {e}")

    # Save the PDF to a BytesIO object
    pdf_output = BytesIO()
    pdf.output(pdf_output)
    pdf_output.seek(0)

    return pdf_output

def generate_resume():
    st.header("Create Your Resume")

    # Taking inputs from the user
    name = st.text_input("Enter Name")
    email = st.text_input("Enter Email")
    phone = st.text_input("Enter Phone Number")
    skills = st.text_area("Enter Skills")
    education = st.text_area("Enter Education")
    work_experience = st.text_area("Enter Work Experience")
    projects = st.text_area("Enter Projects")
    achievements = st.text_area("Enter Achievements")
    certifications = st.text_area("Enter Certifications")
    hobbies = st.text_area("Enter Hobbies")

    # Once all inputs are received
    if st.button("Generate Resume"):
        # Prepare the input prompt for AI generation
        input_prompt9 = f"""
        You are a skilled ATS (Applicant Tracking System) scanner with a deep understanding of data science and ATS functionality.
        Your task is to create the resume for the provided job description and return the resume that was created by using the user inputs:
        Name: {name}
        Email: {email}
        Phone Number: {phone}

        Skills: {skills}

        Education: {education}

        Work Experience: {work_experience}

        Projects: {projects}

        Achievements: {achievements}

        Certifications: {certifications}

        Hobbies: {hobbies}
        The output should come as a PDF file containing the generated resume.
        """

        # Get the response from the AI model
        response = get_gemini_response1(input_prompt9)

        # Display the generated resume to the user
        st.subheader("Generated Resume")
        st.write(response)

        # Generate a PDF from the response
        pdf = generate_pdf(response)

        # Provide a download button for the generated PDF
        st.download_button(
            label="Download Resume as PDF",
            data=pdf,
            file_name=f"{name}_resume.pdf",
            mime="application/pdf"
        )




def tailor_resume(resume_text, job_desc_text):
    # Extract skills from the job description
    jd_skills = extract_skills(job_desc_text)

    # Tokenize and clean the resume text
    resume_tokens = word_tokenize(resume_text.lower())
    resume_cleaned_tokens = [word for word in resume_tokens if word.isalpha() and word not in stopwords.words('english')]

    # Identify relevant and missing skills
    relevant_resume_skills = set(resume_cleaned_tokens).intersection(jd_skills)
    missing_skills = jd_skills - relevant_resume_skills

    # Load resume into a docx template for editing
    doc = Document()
    doc.add_heading('Tailored Resume', 0)

    # Split resume text into paragraphs
    resume_lines = resume_text.split("\n")

    # Add the resume content to the docx, while suggesting improvements
    for line in resume_lines:
        doc.add_paragraph(line)

    # Insert a section for adding missing skills
    if missing_skills:
        doc.add_heading('Suggested Improvements (Based on Job Description)', level=1)
        doc.add_paragraph(
            "The following skills are suggested to be added based on the job description:\n" +
            ', '.join(missing_skills) + ".\nPlease consider adding them to your resume, especially in sections like 'Skills' or 'Experience'."
        )

    # Add relevant skills (already in the resume)
    doc.add_heading('Relevant Skills Already Present', level=1)
    doc.add_paragraph(', '.join(relevant_resume_skills))

    # Return the tailored resume as a downloadable DOCX file
    output = BytesIO()
    doc.save(output)
    return output.getvalue()

# Streamlit UI
st.title("GLA University ATS System")
st.subheader("About")
st.write("This sophisticated ATS project, developed with Gemini Pro and Streamlit, seamlessly incorporates advanced features including resume match percentage, keyword analysis to identify missing criteria, and the generation of comprehensive profile summaries, enhancing the efficiency and precision of the candidate evaluation process for discerning talent acquisition professionals.")

st.markdown("""
  - [Streamlit](https://streamlit.io/)
  - [Gemini Pro](https://deepmind.google/technologies/gemini/#introduction)
  - [makersuit API Key](https://makersuite.google.com/)

  """)

# Sidebar for input
st.sidebar.header("Upload Your Job Description")
job_desc_file = st.sidebar.file_uploader("Upload Job Description (PDF)", type="pdf")
op=st.sidebar.selectbox("Resume:",["Choose an option","Yes, I have","No, I have to create."])

#Prompts
input_prompt1 = """
 You are an experienced Technical Human Resource Manager,your task is to review the provided resume against the job description.
  Please share your professional evaluation on whether the candidate's profile aligns with the role.
 Highlight the strengths and weaknesses of the applicant in relation to the specified job requirements.
"""

input_prompt3 = """
You are an skilled ATS (Applicant Tracking System) scanner with a deep understanding of data science and ATS functionality,
your task is to evaluate the resume against the provided job description. give me the percentage of match if the resume matches
the job description. The output should come as percentage.
"""

input_prompt4 = """
You are an skilled ATS (Applicant Tracking System) scanner with a deep understanding of data science and ATS functionality,
your task is to evaluate the resume against the provided job description. give me the relevant skills if the resume matches
the job description. The output should come as text containing all relevant skills required for given job description .
"""

input_prompt5 = """
You are an skilled ATS (Applicant Tracking System) scanner with a deep understanding of data science and ATS functionality,
your task is to evaluate the resume against the provided job description. give me the non-relevant skills if the resume matches
the job description. The output should come as text containing all non-relevant skills mentioned in resume that are not required for given job description .
"""

input_prompt6 = """
You are a skilled ATS (Applicant Tracking System) scanner with a deep understanding of data science and ATS functionality.
Your task is to evaluate the resume against the provided job description and return only the plagiarism score, expressed
as a percentage of similarity between the two documents.
"""

input_prompt7 = """
You are a skilled ATS (Applicant Tracking System) scanner with a deep understanding of data science and ATS functionality.
Your task is to evaluate the resume against the provided job description and return only the Relevant Projects, for the given job description.
The output should come as text containing all relevant projects required for given job description.
"""

input_prompt8 = """
You are a skilled ATS (Applicant Tracking System) scanner with a deep understanding of data science and ATS functionality.
Your task is to evaluate the resume against the provided job description and return only the Recommended Skills required for Job Description, for the given resume.
The output should come as text containing all recommended skills required for given job description.
"""


if op=="Yes, I have":
  pdf_file = st.sidebar.file_uploader("Upload Resume (PDF)", type="pdf")

  opt = st.sidebar.selectbox("Available Options", ["Choose an option","Percentage match", "Show Relevant Skills", "Non-relevant Skills", "Plagiarism Score", "Relevant Projects", "Recommended Skills", "Tell Me About the Resume" , "Tailor Resume"])

  # Option: Job Description Matching Score
  if opt == "Percentage match":
      if pdf_file is not None and job_desc_file is not None:
          pdf_content = input_pdf_setup(pdf_file)
          job_desc_text = input_pdf_setup(job_desc_file)
          response = get_gemini_response(input_prompt3, pdf_content, job_desc_text[0])
          st.subheader("Percentage Match")
          st.write(response)
      else:
          st.write("Please upload both the resume and job description.")
  # Option: Show Relevant Skills
  if opt == "Show Relevant Skills":
      if pdf_file is not None and job_desc_file is not None:
          pdf_content = input_pdf_setup(pdf_file)
          job_desc_text = input_pdf_setup(job_desc_file)  # Fix: initialize job_desc_text
          response = get_gemini_response(input_prompt4, pdf_content, job_desc_text[0])
          st.subheader("Relevant Skills")
          st.write(response)
      else:
          st.write("Please upload both the resume and job description.")

  # Option: Non-relevant Skills
  if opt == "Non-relevant Skills":
      if pdf_file is not None and job_desc_file is not None:
          pdf_content = input_pdf_setup(pdf_file)
          job_desc_text = input_pdf_setup(job_desc_file)  # Fix: initialize job_desc_text
          response = get_gemini_response(input_prompt5, pdf_content, job_desc_text[0])
          st.subheader("Non-Relevant Skills")
          st.write(response)
      else:
          st.write("Please upload both the resume and job description.")

  # Option: Plagiarism Score
  if opt == "Plagiarism Score":
      if pdf_file is not None and job_desc_file is not None:
          pdf_content = input_pdf_setup(pdf_file)
          job_desc_text = input_pdf_setup(job_desc_file)  # Fix: initialize job_desc_text
          response = get_gemini_response(input_prompt6, pdf_content, job_desc_text[0])
          st.subheader("Plagiarism Score")
          st.write(response)
      else:
          st.write("Please upload both the resume and job description.")

  # Option: Relevant Projects
  if opt == "Relevant Projects":
      if pdf_file is not None and job_desc_file is not None:
          pdf_content = input_pdf_setup(pdf_file)
          job_desc_text = input_pdf_setup(job_desc_file)  # Fix: initialize job_desc_text
          response = get_gemini_response(input_prompt7, pdf_content, job_desc_text[0])
          st.subheader("Relevant Projects for the Job Description")
          st.write(response)
      else:
          st.write("Please upload both the resume and job description.")


  # Option: Recommended Skills
  if opt == "Recommended Skills":
      if pdf_file is not None and job_desc_file is not None:
          pdf_content = input_pdf_setup(pdf_file)
          job_desc_text = input_pdf_setup(job_desc_file)  # Fix: initialize job_desc_text
          response = get_gemini_response(input_prompt8, pdf_content, job_desc_text[0])
          st.subheader("Recommended Skills")
          st.write(response)
      else:
          st.write("Please upload both the resume and job description.")

  # Option: Resume Summary
  if opt == "Tell Me About the Resume":
      if pdf_file is not None and job_desc_file is not None:
          pdf_content = input_pdf_setup(pdf_file)
          job_desc_text = input_pdf_setup(job_desc_file)  # Fix: initialize job_desc_text
          response = get_gemini_response(input_prompt1, pdf_content, job_desc_text[0])
          st.subheader("Resume Tells")
          st.write(response)
      else:
          st.write("Please upload both the resume and job description.")

  # Option: Tailor Resume
  if opt == "Tailor Resume":
      if pdf_file is not None and job_desc_file is not None:
          pdf_content = input_pdf_setup(pdf_file)
          job_desc_text = input_pdf_setup(job_desc_file)
          tailored_resume = tailor_resume(pdf_content[0], job_desc_text[0])

          st.download_button(
              label="Download Tailored Resume (DOCX)",
              data=tailored_resume,
              file_name="tailored_resume.docx",
              mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
          )

          st.header("Tailored Resume Suggestions")
          st.write("Your resume has been tailored with suggestions on adding missing skills and highlighting relevant ones.")
      else:
          st.write("Please upload both the resume and job description.")

if op=="No, I have to create.":
  generate_resume()
